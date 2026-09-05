# Minimal runnable fixtures for the snippets called out in review: one file per format.
# Each script is self-contained, writes only scratch files into the current directory,
# and exits non-zero on failed assertions. Run from any scratch directory:
#   python pdf_fixture.py   (deps: reportlab, pypdf, pymupdf)
#   python pptx_fixture.py  (deps: python-pptx)
#   python xlsx_fixture.py  (deps: openpyxl)
#   python docx_fixture.py  (deps: python-docx, pymupdf, soffice on PATH)
import copy
import os
import subprocess
import sys
import unicodedata
import zipfile
from contextlib import contextmanager
from pathlib import Path, PurePosixPath
from tempfile import TemporaryDirectory, TemporaryFile

import fitz
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

failures = []


def check(name, cond, extra=""):
    print(("PASS " if cond else "FAIL ") + name + ((" :: " + str(extra)) if not cond and extra else ""))
    if not cond:
        failures.append(name)


# ---- review.md health check: external entities stay unresolved ---------------
safe_xml_parser = etree.XMLParser(
    load_dtd=False,
    resolve_entities=False,
    no_network=True,
    huge_tree=False,
    recover=False,
)
hostile_xml = (
    b'<!DOCTYPE root [<!ENTITY xxe SYSTEM "file:///definitely-not-readable">]>'
    b'<root>&xxe;</root>'
)
parsed = etree.fromstring(hostile_xml, parser=safe_xml_parser)
check("DOCX XML parser leaves external entities unresolved", parsed.text is None and len(parsed) == 1)

# ---- review.md health check rejects archive bombs before expanding parts -------
MAX_ARCHIVE_BYTES = 200 * 1024 * 1024
MAX_MEMBERS = 10_000
MAX_XML_PART = 20 * 1024 * 1024
MAX_ENTRY = 100 * 1024 * 1024
MAX_TOTAL_UNCOMPRESSED = 500 * 1024 * 1024
MAX_COMPRESSION_RATIO = 200
CONTENT_TYPES_NAMESPACE = "http://schemas.openxmlformats.org/package/2006/content-types"
MAX_MEMBER_COMPONENT_BYTES = 255
MAX_MEMBER_COMPONENT_UTF16_UNITS = 255
MAX_MEMBER_PATH_BYTES = 1024
MAX_MEMBER_PATH_UTF16_UNITS = 240
MAX_MEMBER_COMPONENTS = 64


def require(condition, message):
    if not condition:
        raise ValueError(message)


def xml_content_type(value):
    media_type = (value or "").split(";", 1)[0].strip().casefold()
    return media_type in {"application/xml", "text/xml"} or media_type.endswith("+xml")


def declared_xml_parts(archive, infos):
    by_name = {info.filename: info for info in infos}
    content_types_info = by_name.get("[Content_Types].xml")
    require(content_types_info is not None, "missing [Content_Types].xml")
    require(content_types_info.file_size <= MAX_XML_PART,
            "oversized XML part: [Content_Types].xml")
    with archive.open(content_types_info) as stream:
        content_types_blob = stream.read(MAX_XML_PART + 1)
    require(len(content_types_blob) <= MAX_XML_PART,
            "oversized XML part: [Content_Types].xml")
    root = etree.fromstring(content_types_blob, parser=safe_xml_parser)
    require(root.tag == f"{{{CONTENT_TYPES_NAMESPACE}}}Types",
            "invalid [Content_Types].xml root")
    defaults = {}
    overrides = {}
    for child in root:
        if child.tag == f"{{{CONTENT_TYPES_NAMESPACE}}}Default":
            extension = (child.get("Extension") or "").casefold()
            require(extension and extension not in defaults,
                    "invalid duplicate content-type default")
            defaults[extension] = child.get("ContentType") or ""
        elif child.tag == f"{{{CONTENT_TYPES_NAMESPACE}}}Override":
            part_name = child.get("PartName") or ""
            require(part_name.startswith("/") and part_name[1:] not in overrides,
                    "invalid duplicate content-type override")
            overrides[part_name[1:]] = child.get("ContentType") or ""
    xml_names = {"[Content_Types].xml"}
    for info in infos:
        suffix = info.filename.rsplit(".", 1)[1].casefold() if "." in info.filename else ""
        content_type = overrides.get(info.filename, defaults.get(suffix, ""))
        if (info.filename.casefold().endswith((".xml", ".rels"))
                or xml_content_type(content_type)):
            xml_names.add(info.filename)
    return xml_names


def validate_docx_package(path):
    require(
        Path(path).stat().st_size <= MAX_ARCHIVE_BYTES,
        "compressed DOCX file size above limit",
    )
    with zipfile.ZipFile(path) as archive:
        infos = archive.infolist()
        require(len(infos) <= MAX_MEMBERS, "archive member count above limit")
        names = {info.filename for info in infos}
        require(len(names) == len(infos), "duplicate archive member names are unsafe")
        require(
            "[Content_Types].xml" in names and "word/document.xml" in names,
            "required DOCX package parts are missing",
        )
        require(
            sum(info.file_size for info in infos) <= MAX_TOTAL_UNCOMPRESSED,
            "declared archive size exceeds the review limit",
        )
        xml_names = declared_xml_parts(archive, infos)
        actual_total = 0
        for info in infos:
            require(info.file_size <= MAX_ENTRY, f"oversized part: {info.filename}")
            require(
                info.file_size / max(info.compress_size, 1) <= MAX_COMPRESSION_RATIO,
                f"suspicious compression ratio: {info.filename}",
            )
            is_xml = info.filename in xml_names
            if is_xml:
                require(info.file_size <= MAX_XML_PART, f"oversized XML part: {info.filename}")
            chunks = []
            actual_size = 0
            with archive.open(info) as stream:
                while chunk := stream.read(64 * 1024):
                    actual_size += len(chunk)
                    actual_total += len(chunk)
                    require(actual_size <= MAX_ENTRY, f"part exceeded read limit: {info.filename}")
                    require(
                        actual_total <= MAX_TOTAL_UNCOMPRESSED,
                        "archive exceeded total read limit",
                    )
                    if is_xml:
                        chunks.append(chunk)
            require(actual_size == info.file_size, f"size mismatch: {info.filename}")
            if is_xml:
                etree.fromstring(b"".join(chunks), parser=safe_xml_parser)


@contextmanager
def validated_docx_source(path):
    with Path(path).open("rb") as external_source, TemporaryFile() as source:
        copied = 0
        while chunk := external_source.read(64 * 1024):
            copied += len(chunk)
            require(copied <= MAX_ARCHIVE_BYTES, "compressed DOCX file size above limit")
            source.write(chunk)
        source.seek(0)
        with zipfile.ZipFile(source) as archive:
            infos = archive.infolist()
            require(len(infos) <= MAX_MEMBERS, "archive member count above limit")
            names = {info.filename for info in infos}
            require(len(names) == len(infos), "duplicate archive member names are unsafe")
            require("[Content_Types].xml" in names and "word/document.xml" in names,
                    "missing required OPC members")
            require(sum(info.file_size for info in infos) <= MAX_TOTAL_UNCOMPRESSED,
                    "declared total uncompressed size above limit")
            xml_names = declared_xml_parts(archive, infos)
            actual_total = 0
            for info in infos:
                require(info.file_size <= MAX_ENTRY, f"oversized part: {info.filename}")
                require(info.file_size / max(info.compress_size, 1) <= MAX_COMPRESSION_RATIO,
                        f"suspicious compression ratio: {info.filename}")
                is_xml = info.filename in xml_names
                if is_xml:
                    require(info.file_size <= MAX_XML_PART,
                            f"oversized XML part: {info.filename}")
                chunks = []
                actual_size = 0
                with archive.open(info) as stream:
                    while chunk := stream.read(64 * 1024):
                        actual_size += len(chunk)
                        actual_total += len(chunk)
                        require(actual_size <= MAX_ENTRY,
                                f"part exceeded read limit: {info.filename}")
                        require(actual_total <= MAX_TOTAL_UNCOMPRESSED,
                                "archive exceeded total read limit")
                        if is_xml:
                            chunks.append(chunk)
                require(actual_size == info.file_size, f"size mismatch: {info.filename}")
                if is_xml:
                    etree.fromstring(b"".join(chunks), parser=safe_xml_parser)
        source.seek(0)
        yield source


def load_validated_docx(path, loader=Document):
    with validated_docx_source(path) as source:
        return loader(source)


health_doc = Document()
health_doc.add_paragraph("bounded health check")
health_doc.save("healthy.docx")
try:
    validate_docx_package("healthy.docx")
    healthy_package_passed = True
except Exception:
    healthy_package_passed = False
check("bounded package health check accepts an ordinary DOCX", healthy_package_passed)

with zipfile.ZipFile("compressed-bomb.docx", "w", zipfile.ZIP_DEFLATED) as archive:
    archive.writestr("[Content_Types].xml", "<Types/>")
    archive.writestr("word/document.xml", "<document>" + (" " * 2_000_000) + "</document>")
try:
    validate_docx_package("compressed-bomb.docx")
    archive_bomb_rejected = False
except (AssertionError, ValueError):
    archive_bomb_rejected = True
check("suspicious compression ratio is rejected before XML expansion", archive_bomb_rejected)
check(
    "archive safety checks remain active under optimized Python",
    __debug__ or archive_bomb_rejected,
)
loader_calls = []
try:
    load_validated_docx(
        "compressed-bomb.docx",
        loader=lambda source: loader_calls.append(source),
    )
    read_bomb_rejected = False
except ValueError:
    read_bomb_rejected = True
check("structured DOCX read rejects a bomb before python-docx is called",
      read_bomb_rejected and loader_calls == [], loader_calls)
loaded_healthy = load_validated_docx("healthy.docx")
check("structured DOCX read validates and loads one private snapshot",
      loaded_healthy.paragraphs[0].text == "bounded health check")


def docx_with_extra_part(source_path, output_path, name, payload, content_type=None):
    with zipfile.ZipFile(source_path) as source:
        members = {info.filename: source.read(info) for info in source.infolist()}
    if content_type is not None:
        root = etree.fromstring(members["[Content_Types].xml"], parser=safe_xml_parser)
        override = etree.SubElement(root, f"{{{CONTENT_TYPES_NAMESPACE}}}Override")
        override.set("PartName", "/" + name)
        override.set("ContentType", content_type)
        members["[Content_Types].xml"] = etree.tostring(root, xml_declaration=True,
                                                         encoding="UTF-8")
    members[name] = payload
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as output:
        for member_name, member_data in members.items():
            output.writestr(member_name, member_data)


docx_with_extra_part("healthy.docx", "uppercase-xml.docx", "custom/BROKEN.XML", b"<broken")
docx_with_extra_part(
    "healthy.docx", "declared-xml.docx", "custom/broken.dat", b"<broken",
    "application/vnd.example.custom+xml",
)
for label, package_path in (
    ("uppercase XML suffix", "uppercase-xml.docx"),
    ("content-type XML override", "declared-xml.docx"),
):
    try:
        validate_docx_package(package_path)
        malformed_declared_xml_rejected = False
    except etree.XMLSyntaxError:
        malformed_declared_xml_rejected = True
    check(f"health check parses {label} parts", malformed_declared_xml_rejected)

with zipfile.ZipFile("many-members.docx", "w", zipfile.ZIP_STORED) as archive:
    archive.writestr("[Content_Types].xml", "<Types/>")
    archive.writestr("word/document.xml", "<document/>")
    for member_index in range(MAX_MEMBERS - 1):
        archive.writestr(f"word/zero-{member_index:05d}.bin", b"")
try:
    validate_docx_package("many-members.docx")
    many_members_rejected = False
except ValueError as exc:
    many_members_rejected = str(exc) == "archive member count above limit"
check(
    "member-count gate rejects 10,001 distinct zero-byte archive members before traversal",
    many_members_rejected,
)

# ---- read.md includes block/inline content controls and controlled tables -------
MC_NAMESPACE = "http://schemas.openxmlformats.org/markup-compatibility/2006"
MC_ALTERNATE_CONTENT = f"{{{MC_NAMESPACE}}}AlternateContent"
MC_CHOICE = f"{{{MC_NAMESPACE}}}Choice"
MC_FALLBACK = f"{{{MC_NAMESPACE}}}Fallback"
SUPPORTED_MC_NAMESPACE_URIS = {
    MC_NAMESPACE,
    qn("w:p").split("}", 1)[0][1:],
    qn("r:id").split("}", 1)[0][1:],
}


def alternate_content_branch(element):
    fallback = None
    for child in element.iterchildren():
        if child.tag == MC_CHOICE:
            required_prefixes = (child.get("Requires") or "").split()
            if required_prefixes and all(
                child.nsmap.get(prefix) in SUPPORTED_MC_NAMESPACE_URIS
                for prefix in required_prefixes
            ):
                return child
        elif child.tag == MC_FALLBACK and fallback is None:
            fallback = child
    return fallback


def unresolved_alternate_content(element):
    return {
        "kind": "AlternateContent",
        "requires": [
            (child.get("Requires") or "").split()
            for child in element.iterchildren()
            if child.tag == MC_CHOICE
        ],
        "reason": "no supported Choice and no Fallback",
    }


def iter_effective_children(root):
    for child in root.iterchildren():
        if child.tag != MC_ALTERNATE_CONTENT:
            yield "element", child
            continue
        branch = alternate_content_branch(child)
        if branch is None:
            yield "unreadable", unresolved_alternate_content(child)
        else:
            yield from iter_effective_children(branch)


def iter_part_blocks(root, parent):
    for child_kind, child in iter_effective_children(root):
        if child_kind == "unreadable":
            yield child_kind, child
            continue
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, parent)
        elif child.tag == qn("w:altChunk"):
            relationship_id = child.get(qn("r:id"))
            part = getattr(parent, "part", None)
            relationship = None if part is None else part.rels.get(relationship_id)
            yield "unreadable", {
                "kind": "altChunk",
                "relationship_id": relationship_id,
                "target": None if relationship is None else relationship.target_ref,
                "content_type": None if relationship is None or relationship.is_external
                else relationship.target_part.content_type,
            }
        else:
            yield from iter_part_blocks(child, parent)


def iter_paragraph_items(paragraph):
    def walk(element):
        for child_kind, child in iter_effective_children(element):
            if child_kind == "unreadable":
                yield child_kind, child
                continue
            if child.tag == qn("w:r"):
                yield "run", Run(child, paragraph)
            elif child.tag != qn("w:p"):
                yield from walk(child)
    yield from walk(paragraph._p)


def iter_paragraph_runs(paragraph):
    for kind, item in iter_paragraph_items(paragraph):
        if kind == "run":
            yield item


def legacy_symbol_record(symbol):
    font = symbol.get(qn("w:font"))
    character = symbol.get(qn("w:char"))
    return f"[unreadable legacy symbol font={font!r} char={character!r}]"


def run_text(run):
    pieces = []
    text_tags = {
        qn("w:br"), qn("w:cr"), qn("w:noBreakHyphen"), qn("w:ptab"),
        qn("w:t"), qn("w:tab"), qn("w:sym"),
    }
    for child_kind, child in iter_effective_children(run._r):
        if child_kind == "unreadable":
            pieces.append(f"[unreadable {child['kind']}: {child['reason']}]")
        elif child.tag in text_tags:
            pieces.append(legacy_symbol_record(child) if child.tag == qn("w:sym") else str(child))
    return "".join(pieces)


def paragraph_text(paragraph):
    pieces = []
    for kind, item in iter_paragraph_items(paragraph):
        pieces.append(
            run_text(item) if kind == "run"
            else f"[unreadable {item['kind']}: {item['reason']}]"
        )
    return "".join(pieces)


def cell_paragraph_text(paragraph):
    pieces = []
    for kind, item in iter_paragraph_items(paragraph):
        if kind == "unreadable":
            pieces.append(f"[unreadable {item['kind']}: {item['reason']}]")
            continue
        for child_kind, child in iter_effective_children(item._r):
            if child_kind == "unreadable":
                pieces.append(f"[unreadable {child['kind']}: {child['reason']}]")
            elif child.tag == qn("w:t"):
                pieces.append(child.text or "")
            elif child.tag in (qn("w:tab"), qn("w:ptab")):
                pieces.append("<tab>")
            elif child.tag in (qn("w:br"), qn("w:cr")):
                pieces.append("<br>")
            elif child.tag == qn("w:noBreakHyphen"):
                pieces.append("-")
            elif child.tag == qn("w:sym"):
                pieces.append(legacy_symbol_record(child))
    return "".join(pieces)


def tc_text(tc, parent):
    paragraphs = []
    for kind, block in iter_part_blocks(tc, parent):
        if kind == "paragraph":
            paragraphs.append(cell_paragraph_text(block))
    return " / ".join(paragraphs)


def iter_content_control_children(root, target_tag):
    for child_kind, child in iter_effective_children(root):
        if child_kind == "unreadable":
            raise ValueError(
                f"unresolved AlternateContent while locating {target_tag}: {child}"
            )
        if child.tag == target_tag:
            yield child
        elif child.tag == qn("w:sdt"):
            for content in child.findall(qn("w:sdtContent")):
                yield from iter_content_control_children(content, target_tag)


def table_content(table):
    rows = []
    for row_element in iter_content_control_children(table._tbl, qn("w:tr")):
        rendered_cells = []
        row_properties = row_element.find(qn("w:trPr"))
        grid_before_node = None if row_properties is None else row_properties.find(qn("w:gridBefore"))
        grid_after_node = None if row_properties is None else row_properties.find(qn("w:gridAfter"))
        grid_before = 0 if grid_before_node is None else int(grid_before_node.get(qn("w:val"), "0"))
        grid_after = 0 if grid_after_node is None else int(grid_after_node.get(qn("w:val"), "0"))
        column = grid_before
        for cell_element in iter_content_control_children(row_element, qn("w:tc")):
            cell_properties = cell_element.find(qn("w:tcPr"))
            grid_span = None if cell_properties is None else cell_properties.find(qn("w:gridSpan"))
            colspan = 1 if grid_span is None else int(grid_span.get(qn("w:val"), "1"))
            vertical = None if cell_properties is None else cell_properties.find(qn("w:vMerge"))
            vertical_merge = None if vertical is None else vertical.get(qn("w:val"), "continue")
            nested_tables = []
            unreadable = []
            for kind, block in iter_part_blocks(cell_element, table):
                if kind == "table":
                    nested_tables.append(table_content(block))
                elif kind == "unreadable":
                    unreadable.append(block)
            rendered_cells.append({
                "column": column, "colspan": colspan,
                "vMerge": vertical_merge,
                "text": tc_text(cell_element, table),
                "tables": nested_tables,
                "unreadable": unreadable,
            })
            column += colspan
        rows.append({
            "grid_before": grid_before, "cells": rendered_cells, "grid_after": grid_after,
        })
    return rows


def wrap_in_sdt(element):
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    element.getparent().replace(element, sdt)
    content.append(element)
    sdt.append(content)


sdt_doc = Document()
sdt_doc.add_paragraph("direct paragraph")
sdt_paragraph = sdt_doc.add_paragraph("inside content control")
wrap_in_sdt(sdt_paragraph._p)
controlled_table = sdt_doc.add_table(rows=1, cols=1)
controlled_table.cell(0, 0).text = "table한"
row_properties = controlled_table.rows[0]._tr.get_or_add_trPr()
grid_before = OxmlElement("w:gridBefore")
grid_before.set(qn("w:val"), "1")
row_properties.append(grid_before)
grid_after = OxmlElement("w:gridAfter")
grid_after.set(qn("w:val"), "2")
row_properties.append(grid_after)
nested_table = controlled_table.cell(0, 0).add_table(rows=1, cols=1)
nested_table.cell(0, 0).text = "nested한"
inline_paragraph = sdt_doc.add_paragraph("before-")
inline_run = inline_paragraph.add_run("inline한")
inline_paragraph.add_run("-after")
wrap_in_sdt(inline_run._r)
inline_paragraph.add_run("-legacy-")
legacy_symbol_run = inline_paragraph.add_run()
legacy_symbol = OxmlElement("w:sym")
legacy_symbol.set(qn("w:font"), "Wingdings")
legacy_symbol.set(qn("w:char"), "F052")
legacy_symbol_run._r.append(legacy_symbol)
inline_paragraph.add_run("-visible")
chunk_part = Part(
    PackURI("/word/altChunk1.html"), "text/html",
    b"<html><body>IMPORTED ALTCHUNK TEXT</body></html>", sdt_doc.part.package,
)
chunk_relationship = sdt_doc.part.relate_to(chunk_part, RT.A_F_CHUNK)
alt_chunk = OxmlElement("w:altChunk")
alt_chunk.set(qn("r:id"), chunk_relationship)
sdt_doc.element.body.insert(len(sdt_doc.element.body) - 1, alt_chunk)
cell_alt_chunk = OxmlElement("w:altChunk")
cell_alt_chunk.set(qn("r:id"), chunk_relationship)
controlled_table.cell(0, 0)._tc.insert(
    len(controlled_table.cell(0, 0)._tc) - 1, cell_alt_chunk
)
controlled_row_element = controlled_table.rows[0]._tr
controlled_cell_element = controlled_row_element.tc_lst[0]
wrap_in_sdt(controlled_cell_element)
wrap_in_sdt(controlled_row_element)
wrap_in_sdt(controlled_table._tbl)
sdt_doc.save("content-control.docx")
sdt_reopened = Document("content-control.docx")
check("doc.paragraphs omits block content-control text (negative control)",
      "inside content control" not in [paragraph.text for paragraph in sdt_reopened.paragraphs])
walked_blocks = list(iter_part_blocks(sdt_reopened.element.body, sdt_reopened))
check("doc.tables omits a table wrapped by a block content control (negative control)",
      len(sdt_reopened.tables) == 0)
check("block walker preserves document order across content controls",
      [kind for kind, _ in walked_blocks]
      == ["paragraph", "paragraph", "table", "paragraph", "unreadable"], walked_blocks)
walked_paragraphs = [block for kind, block in walked_blocks if kind == "paragraph"]
walked_text = [paragraph_text(paragraph) for paragraph in walked_paragraphs]
check("content-control traversal emits the nested paragraph", "inside content control" in walked_text, walked_text)
check("content-control traversal emits inline run text",
      any("before-inline한-after" in text for text in walked_text), walked_text)
legacy_symbol_marker = "[unreadable legacy symbol font='Wingdings' char='F052']"
legacy_paragraph = next(paragraph for paragraph in walked_paragraphs
                        if "-legacy-" in paragraph_text(paragraph))
check("Run.text silently omits visible legacy w:sym content (negative control)",
      legacy_symbol_marker not in "".join(run.text for run in iter_paragraph_runs(legacy_paragraph)))
check("paragraph extraction preserves the legacy symbol's position, font, and character code",
      f"-legacy-{legacy_symbol_marker}-visible" in paragraph_text(legacy_paragraph),
      paragraph_text(legacy_paragraph))
walked_tables = [block for kind, block in walked_blocks if kind == "table"]
check("Table.rows omits a row wrapped by row-level sdtContent (negative control)",
      len(walked_tables) == 1 and len(walked_tables[0].rows) == 0)
physical_rows = list(iter_content_control_children(walked_tables[0]._tbl, qn("w:tr")))
check("tr.tc_lst omits a cell wrapped by cell-level sdtContent (negative control)",
      len(physical_rows) == 1 and len(physical_rows[0].tc_lst) == 0)
rendered_tables = [table_content(table) for table in walked_tables]
check("content-control traversal emits a wrapped table",
      len(walked_tables) == 1 and "table한" in str(rendered_tables), rendered_tables)
check("table traversal emits the sdt-wrapped physical row and cell",
      len(rendered_tables[0]) == 1 and len(rendered_tables[0][0]["cells"]) == 1,
      rendered_tables)
all_emitted = walked_text + [str(table) for table in rendered_tables]
check("table text is emitted exactly once, not again as prose",
      sum(item.count("table한") for item in all_emitted) == 1, all_emitted)
check("nested-table text is emitted exactly once",
      sum(item.count("nested한") for item in all_emitted) == 1, all_emitted)
check("nonuniform table rows preserve leading/trailing grid omissions",
      rendered_tables[0][0]["grid_before"] == 1
      and rendered_tables[0][0]["cells"][0]["column"] == 1
      and rendered_tables[0][0]["grid_after"] == 2,
      rendered_tables[0])
unreadable_parts = [block for kind, block in walked_blocks if kind == "unreadable"]
expected_alt_chunk = {
    "kind": "altChunk",
    "relationship_id": chunk_relationship,
    "target": "altChunk1.html",
    "content_type": "text/html",
}
check("altChunk content is reported instead of silently omitted",
      unreadable_parts == [expected_alt_chunk], unreadable_parts)
cell_unreadable = rendered_tables[0][0]["cells"][0]["unreadable"]
check("altChunk content inside a table cell is also reported",
      cell_unreadable == [expected_alt_chunk], rendered_tables)
check("altChunk payload is not misrepresented as extracted paragraph text",
      all("IMPORTED ALTCHUNK TEXT" not in text for text in walked_text), walked_text)


# ---- read.md chooses exactly one mc:AlternateContent branch ---------------------
W14_NAMESPACE = "http://schemas.microsoft.com/office/word/2010/wordml"
W_NAMESPACE = qn("w:p").split("}", 1)[0][1:]


def wml_run(text):
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    return run


def wml_paragraph(text):
    paragraph = OxmlElement("w:p")
    paragraph.append(wml_run(text))
    return paragraph


def wml_table(text):
    source = Document()
    table = source.add_table(rows=1, cols=1)
    table.cell(0, 0).text = text
    return copy.deepcopy(table._tbl)


def alternate_content(requires, choice_elements, fallback_elements=None):
    element = etree.Element(
        MC_ALTERNATE_CONTENT,
        nsmap={"mc": MC_NAMESPACE, "w14": W14_NAMESPACE, "w": W_NAMESPACE},
    )
    choice = etree.SubElement(element, MC_CHOICE)
    choice.set("Requires", requires)
    for child in choice_elements:
        choice.append(child)
    if fallback_elements is not None:
        fallback = etree.SubElement(element, MC_FALLBACK)
        for child in fallback_elements:
            fallback.append(child)
    return element


mc_doc = Document()
mc_doc.element.body.insert(
    len(mc_doc.element.body) - 1,
    alternate_content(
        "w14",
        [wml_paragraph("UNSUPPORTED-CHOICE-BLOCK"), wml_table("UNSUPPORTED-CHOICE-TABLE")],
        [wml_paragraph("FALLBACK-BLOCK"), wml_table("FALLBACK-TABLE")],
    ),
)
mc_doc.element.body.insert(
    len(mc_doc.element.body) - 1,
    alternate_content(
        "w",
        [wml_paragraph("SUPPORTED-CHOICE-BLOCK")],
        [wml_paragraph("UNSELECTED-FALLBACK-BLOCK")],
    ),
)
mc_doc.element.body.insert(
    len(mc_doc.element.body) - 1,
    alternate_content("w14", [wml_paragraph("UNRESOLVED-HIDDEN-BLOCK")]),
)
mc_inline = mc_doc.add_paragraph("INLINE-BEFORE-")
mc_inline._p.append(
    alternate_content(
        "w14", [wml_run("UNSUPPORTED-CHOICE-INLINE")], [wml_run("FALLBACK-INLINE")]
    )
)
mc_inline.add_run("-INLINE-AFTER")
mc_doc.save("alternate-content.docx")
mc_reopened = Document("alternate-content.docx")
mc_blocks = list(iter_part_blocks(mc_reopened.element.body, mc_reopened))
mc_paragraph_texts = [paragraph_text(block) for kind, block in mc_blocks if kind == "paragraph"]
mc_tables = [table_content(block) for kind, block in mc_blocks if kind == "table"]
mc_unreadable = [block for kind, block in mc_blocks if kind == "unreadable"]
mc_emitted = mc_paragraph_texts + [str(table) for table in mc_tables]
check("unsupported AlternateContent Choice selects only its Fallback paragraph and table",
      any("FALLBACK-BLOCK" in text for text in mc_emitted)
      and any("FALLBACK-TABLE" in text for text in mc_emitted)
      and all("UNSUPPORTED-CHOICE" not in text for text in mc_emitted),
      mc_emitted)
check("supported AlternateContent Choice wins over its Fallback",
      any("SUPPORTED-CHOICE-BLOCK" in text for text in mc_emitted)
      and all("UNSELECTED-FALLBACK-BLOCK" not in text for text in mc_emitted),
      mc_emitted)
check("inline AlternateContent contributes exactly one selected run",
      "INLINE-BEFORE-FALLBACK-INLINE-INLINE-AFTER" in mc_paragraph_texts
      and all("UNSUPPORTED-CHOICE-INLINE" not in text for text in mc_paragraph_texts),
      mc_paragraph_texts)
check("AlternateContent without a usable branch is explicitly unreadable",
      mc_unreadable == [{
          "kind": "AlternateContent",
          "requires": [["w14"]],
          "reason": "no supported Choice and no Fallback",
      }], mc_unreadable)
check("unresolved AlternateContent payload is not misrepresented as extracted text",
      all("UNRESOLVED-HIDDEN-BLOCK" not in text for text in mc_emitted), mc_emitted)

# Per-run glyph validation must not let a different referenced font hide a missing glyph.
fixture_cmaps = {"CJK Face": {ord("漢")}, "Latin Face": {ord("A")}}
assigned_runs = [("CJK Face", "漢"), ("Latin Face", "漢")]
pooled_passes = all(any(ord(ch) in cmap for cmap in fixture_cmaps.values())
                    for _, text in assigned_runs for ch in text)
per_run_missing = [(face, ch) for face, text in assigned_runs for ch in text
                   if ord(ch) not in fixture_cmaps[face]]
check("pooled cmap is proven unsafe (negative control)", pooled_passes)
check("per-run cmap check identifies the actual missing glyph",
      per_run_missing == [("Latin Face", "漢")], per_run_missing)


def require_clean_cjk_glyph_audit(unresolved, missing):
    if unresolved:
        raise ValueError(f"font files not resolved per run: {unresolved}")
    if missing:
        raise ValueError(f"glyph missing from the run's effective font: {missing}")


try:
    require_clean_cjk_glyph_audit([(0, "漢", "Unresolved Face")], [])
    unresolved_face_rejected = False
except ValueError:
    unresolved_face_rejected = True
check(
    "mandatory CJK audit rejects unresolved font files under optimized Python",
    unresolved_face_rejected,
)

try:
    require_clean_cjk_glyph_audit([], [(0, "漢", "Missing Glyph Face")])
    missing_glyph_rejected = False
except ValueError:
    missing_glyph_rejected = True
check(
    "mandatory CJK audit rejects missing glyphs under optimized Python",
    missing_glyph_rejected,
)

def font_slot(character):
    codepoint = ord(character)
    uses_east_asian_slot = (
        0x1100 <= codepoint <= 0x11FF
        or 0x2F00 <= codepoint <= 0x9FFF
        or 0xA000 <= codepoint <= 0xA4CF
        or 0xA960 <= codepoint <= 0xA97F
        or 0xAC00 <= codepoint <= 0xD7FF
        or 0xF900 <= codepoint <= 0xFAFF
        or 0xFE30 <= codepoint <= 0xFE6F
        or 0xFF00 <= codepoint <= 0xFFEF
        or 0x20000 <= codepoint <= 0x3134F
        or 0x31350 <= codepoint <= 0x33479
    )
    return "eastAsia" if uses_east_asian_slot else ("ascii" if codepoint < 128 else "hAnsi")


check("Han ideographs use the east-Asian font slot", font_slot("漢") == "eastAsia")
check("Hangul syllables use the east-Asian font slot", font_slot("한") == "eastAsia")
check("Hangul Jamo use the east-Asian font slot", font_slot("ᄒ") == "eastAsia")
check("Yi syllables use the east-Asian font slot", font_slot("ꀀ") == "eastAsia")
check("CJK compatibility forms use the east-Asian font slot", font_slot("︰") == "eastAsia")
check("Unicode 17 Han Extension H start uses the east-Asian font slot",
      font_slot(chr(0x31350)) == "eastAsia")
check("Unicode 17 Han Extension J end uses the east-Asian font slot",
      font_slot(chr(0x33479)) == "eastAsia")
check("codepoint after the Unicode 17 Han ranges stays in hAnsi",
      font_slot(chr(0x3347A)) == "hAnsi")
check("ASCII text keeps the ascii font slot", font_slot("A") == "ascii")

# Header/footer parts have no `.document`; effective styles close over the document.
font_doc = Document()
normal_style = font_doc.styles["Normal"]
normal_fonts = normal_style.element.get_or_add_rPr().get_or_add_rFonts()
normal_fonts.set(qn("w:ascii"), "Latin Face")
normal_fonts.set(qn("w:hAnsi"), "Latin Face")
normal_fonts.set(qn("w:eastAsia"), "CJK Face")
header_run = font_doc.sections[0].header.paragraphs[0].add_run("页眉")
footer_run = font_doc.sections[0].footer.paragraphs[0].add_run("页脚")


def face_from_rpr(rpr, slot):
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    theme_attribute = {
        "ascii": "asciiTheme", "hAnsi": "hAnsiTheme",
        "eastAsia": "eastAsiaTheme", "cs": "cstheme",
    }[slot]
    theme_token = rfonts.get(qn("w:" + theme_attribute))
    if theme_token is not None:
        literal = rfonts.get(qn("w:" + slot))
        raise LookupError(
            f"unresolved direct {theme_attribute}={theme_token!r}"
            + (f" alongside {slot}={literal!r}" if literal is not None else "")
        )
    return rfonts.get(qn("w:" + slot))


def style_faces(style, slot):
    while style is not None:
        face = face_from_rpr(style.element.find(qn("w:rPr")), slot)
        if face:
            yield face
        style = style.base_style


def table_style_faces(run, slot):
    element = run._r
    while element is not None and element.tag != qn("w:tbl"):
        element = element.getparent()
    if element is None:
        return []
    table_properties = element.find(qn("w:tblPr"))
    table_style = None if table_properties is None else table_properties.find(qn("w:tblStyle"))
    style_id = None if table_style is None else table_style.get(qn("w:val"))
    style = None if not style_id else font_doc.styles.get_by_id(style_id, WD_STYLE_TYPE.TABLE)
    faces = []
    while style is not None:
        rprs = [style.element.find(qn("w:rPr"))]
        rprs.extend(region.find(qn("w:rPr"))
                    for region in style.element.findall(qn("w:tblStylePr")))
        faces.extend(face for rpr in rprs if (face := face_from_rpr(rpr, slot)))
        style = style.base_style
    return list(dict.fromkeys(faces))


def effective_face(run, slot):
    direct = face_from_rpr(run._r.find(qn("w:rPr")), slot)
    if direct:
        return direct
    if faces := table_style_faces(run, slot):
        raise LookupError(f"conditional table style face requires rendered resolution: {faces}")
    for style in (run.style, run._parent.style, normal_style):
        if face := next(style_faces(style, slot), None):
            return face
    raise LookupError(slot)


theme_attribute_by_slot = {
    "ascii": "asciiTheme", "hAnsi": "hAnsiTheme",
    "eastAsia": "eastAsiaTheme", "cs": "cstheme",
}
for slot, theme_attribute in theme_attribute_by_slot.items():
    themed_run = font_doc.add_paragraph().add_run("漢" if slot == "eastAsia" else "A")
    themed_fonts = themed_run._r.get_or_add_rPr().get_or_add_rFonts()
    themed_fonts.set(qn("w:" + theme_attribute), "majorEastAsia" if slot == "eastAsia" else "minorAscii")
    try:
        effective_face(themed_run, slot)
        direct_theme_rejected = False
    except LookupError:
        direct_theme_rejected = True
    check(f"direct {theme_attribute} does not fall through to an inherited literal face",
          direct_theme_rejected)

ambiguous_theme_run = font_doc.add_paragraph().add_run("漢")
ambiguous_fonts = ambiguous_theme_run._r.get_or_add_rPr().get_or_add_rFonts()
ambiguous_fonts.set(qn("w:eastAsia"), "Literal Face")
ambiguous_fonts.set(qn("w:eastAsiaTheme"), "majorEastAsia")
try:
    effective_face(ambiguous_theme_run, "eastAsia")
    literal_and_theme_rejected = False
except LookupError:
    literal_and_theme_rejected = True
check("same-slot literal plus theme font declaration fails closed",
      literal_and_theme_rejected)


check("header part has no document back-reference (negative control)",
      not hasattr(header_run.part, "document"))
check("header runs resolve Normal from the owning document",
      effective_face(header_run, "eastAsia") == "CJK Face")
check("footer runs resolve Normal from the owning document",
      effective_face(footer_run, "eastAsia") == "CJK Face")
inline_paragraph_reopened = next(
    paragraph for paragraph in walked_paragraphs if "before-" in paragraph_text(paragraph)
)
check("inline content-control runs are included in glyph traversal",
      any(run.text == "inline한" for run in iter_paragraph_runs(inline_paragraph_reopened)))

hyperlink_paragraph = font_doc.add_paragraph()
hyperlink = OxmlElement("w:hyperlink")
hyperlink.set(qn("w:anchor"), "fixture-target")
hyperlink_run = OxmlElement("w:r")
hyperlink_text = OxmlElement("w:t")
hyperlink_text.text = "链接漢"
hyperlink_run.append(hyperlink_text)
hyperlink.append(hyperlink_run)
hyperlink_paragraph._p.append(hyperlink)
walked_hyperlink_runs = list(iter_paragraph_runs(hyperlink_paragraph))
check("Paragraph.runs omits hyperlink runs (negative control)",
      all(run.text != "链接漢" for run in hyperlink_paragraph.runs))
check("glyph traversal includes CJK text nested in a hyperlink",
      [run.text for run in walked_hyperlink_runs] == ["链接漢"],
      [run.text for run in walked_hyperlink_runs])
check("hyperlink CJK text resolves through the east-Asian font slot",
      effective_face(walked_hyperlink_runs[0], "eastAsia") == "CJK Face")

conditional_style = font_doc.styles.add_style("Conditional CJK Table", WD_STYLE_TYPE.TABLE)
first_row = OxmlElement("w:tblStylePr")
first_row.set(qn("w:type"), "firstRow")
conditional_rpr = OxmlElement("w:rPr")
conditional_fonts = OxmlElement("w:rFonts")
conditional_fonts.set(qn("w:eastAsia"), "Conditional CJK Face")
conditional_rpr.append(conditional_fonts)
first_row.append(conditional_rpr)
conditional_style.element.append(first_row)
font_table = font_doc.add_table(rows=1, cols=1)
font_table.style = conditional_style
conditional_run = font_table.cell(0, 0).paragraphs[0].add_run("漢")
try:
    effective_face(conditional_run, "eastAsia")
    conditional_style_rejected = False
except LookupError:
    conditional_style_rejected = True
check("glyph audit fails closed for conditional table-style fonts",
      conditional_style_rejected)

# ---- edit.md guarded cross-run replacement ------------------------------------
SAFE_RUN_CHILDREN = {
    qn("w:rPr"), qn("w:t"), qn("w:tab"), qn("w:cr"),
}
MODELED_PARAGRAPH_CHILDREN = {qn("w:pPr"), qn("w:r")}


def unsafe_run_content(run):
    unsafe = []
    for child in run._r:
        ordinary_break = child.tag == qn("w:br") and dict(child.attrib) in (
            {}, {qn("w:type"): "textWrapping"},
        )
        if child.tag not in SAFE_RUN_CHILDREN and not ordinary_break:
            unsafe.append(child.tag.rsplit("}", 1)[-1])
    return unsafe


def replace_across_runs(paragraph, old, new):
    if not old:
        raise ValueError("old must not be empty")
    unmodeled = [
        child.tag.rsplit("}", 1)[-1]
        for child in paragraph._p
        if child.tag not in MODELED_PARAGRAPH_CHILDREN
    ]
    if unmodeled:
        raise ValueError(f"paragraph contains unmodeled inline containers: {unmodeled}")
    runs = list(paragraph.runs)
    text = "".join(run.text for run in runs)
    starts = []
    position = 0
    while (start := text.find(old, position)) != -1:
        starts.append(start)
        position = start + len(old)
    spans = []
    position = 0
    for index, run in enumerate(runs):
        end = position + len(run.text)
        if end > position:
            spans.append((index, position, end))
        position = end
    matches = []
    for start in starts:
        end = start + len(old)
        first, first_start, _ = next(s for s in spans if s[1] <= start < s[2])
        last, last_start, _ = next(s for s in spans if s[1] < end <= s[2])
        matches.append((start, end, first, first_start, last, last_start))
    affected_indexes = {
        index
        for _, _, first, _, last, _ in matches
        for index in range(first, last + 1)
    }
    unsafe = {}
    for index in affected_indexes:
        if children := unsafe_run_content(runs[index]):
            unsafe[index] = children
    if unsafe:
        raise ValueError(f"matched runs contain non-text content: {unsafe}")
    for start, end, first, first_start, last, last_start in reversed(matches):
        prefix = runs[first].text[:start - first_start]
        suffix = runs[last].text[end - last_start:]
        if first == last:
            runs[first].text = prefix + new + suffix
        else:
            runs[first].text = prefix + new
            for index in range(first + 1, last):
                runs[index].text = ""
            runs[last].text = suffix
    return len(starts)


safe_doc = Document()
safe_paragraph = safe_doc.add_paragraph()
safe_first = safe_paragraph.add_run("T")
safe_first.bold = True
safe_paragraph.add_run("B")
safe_paragraph.add_run("D")
check("text-only cross-run match is replaced", replace_across_runs(safe_paragraph, "TBD", "Done") == 1)
check("safe replacement keeps first-run formatting", safe_paragraph.text == "Done" and safe_first.bold)

icon = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 4, 4), False)
icon.clear_with(200)
icon.save("inline-icon.png")
guard_doc = Document()
guard_paragraph = guard_doc.add_paragraph()
guard_run = guard_paragraph.add_run("TBD")
guard_run.add_picture("inline-icon.png")
try:
    replace_across_runs(guard_paragraph, "TBD", "Done")
    rejected_drawing_run = False
except ValueError:
    rejected_drawing_run = True
check("replacement rejects a matched run containing a drawing", rejected_drawing_run)
check(
    "rejected replacement leaves text and drawing untouched",
    guard_run.text == "TBD" and len(guard_run._r.findall(qn("w:drawing"))) == 1,
)

clear_doc = Document()
clear_paragraph = clear_doc.add_paragraph()
clear_run = clear_paragraph.add_run("TBD")
clear_break = OxmlElement("w:br")
clear_break.set(qn("w:type"), "textWrapping")
clear_break.set(qn("w:clear"), "left")
clear_run._r.append(clear_break)
clear_before = etree.tostring(clear_run._r)
try:
    replace_across_runs(clear_paragraph, "TBD", "done")
    clear_break_rejected = False
except ValueError:
    clear_break_rejected = True
check("replacement rejects a wrapping break with clear semantics", clear_break_rejected)
check("rejected clear-break replacement is atomic",
      etree.tostring(clear_run._r) == clear_before)

container_doc = Document()
container_paragraph = container_doc.add_paragraph()
container_paragraph.add_run("T")
container_hyperlink = OxmlElement("w:hyperlink")
container_hyperlink.set(qn("w:anchor"), "fixture-target")
container_link_run = OxmlElement("w:r")
container_link_text = OxmlElement("w:t")
container_link_text.text = "link"
container_link_run.append(container_link_text)
container_hyperlink.append(container_link_run)
container_paragraph._p.append(container_hyperlink)
container_paragraph.add_run("BD")
container_before = etree.tostring(container_paragraph._p)
check("Paragraph.runs can manufacture a false match across a hyperlink (negative control)",
      "".join(run.text for run in container_paragraph.runs) == "TBD"
      and paragraph_text(container_paragraph) == "TlinkBD")
try:
    replace_across_runs(container_paragraph, "TBD", "Done")
    rejected_inline_container = False
except ValueError:
    rejected_inline_container = True
check("replacement rejects unmodeled inline containers before matching",
      rejected_inline_container)
check("rejected inline-container replacement is atomic",
      etree.tostring(container_paragraph._p) == container_before
      and paragraph_text(container_paragraph) == "TlinkBD")


def list_number_num_id(doc):
    """The numId that the ListNumber style binds to in this document part."""
    styles = doc.part.element.body.getparent()  # document.xml root; styles live in another part
    styles_part = doc.part.part_related_by(RT.STYLES)
    for style in styles_part.element.findall(qn("w:style")):
        if style.get(qn("w:styleId")) == "ListNumber":
            numPr = style.find(qn("w:pPr") + "/" + qn("w:numPr"))
            if numPr is not None:
                return int(numPr.find(qn("w:numId")).get(qn("w:val")))
    raise LookupError("ListNumber style has no numPr")


def new_restart_num_id(doc, base_num_id):
    """Clone <w:num> base_num_id with a startOverride so the next list restarts at 1."""
    numbering = doc.part.part_related_by(RT.NUMBERING).element
    source = next(
        n for n in numbering.findall(qn("w:num"))
        if n.get(qn("w:numId")) == str(base_num_id)
    )
    clone = copy.deepcopy(source)
    new_id = max(int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))) + 1
    clone.set(qn("w:numId"), str(new_id))
    level_zero_overrides = [
        item for item in clone.findall(qn("w:lvlOverride"))
        if item.get(qn("w:ilvl")) == "0"
    ]
    if len(level_zero_overrides) > 1:
        raise ValueError("base numbering has duplicate level-zero overrides")
    if level_zero_overrides:
        override = level_zero_overrides[0]
        for old_start in override.findall(qn("w:startOverride")):
            override.remove(old_start)
    else:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        clone.append(override)
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.insert(0, start)
    numbering.append(clone)
    return new_id


def numbered_paragraph(doc, text, num_id):
    p = doc.add_paragraph(text, style="List Number")
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    return p


# ---- document A: naive reuse of List Number (second list continues) -----------
doc_a = Document()
doc_a.add_paragraph("First list")
for item in ("one", "two", "three"):
    doc_a.add_paragraph(item, style="List Number")
doc_a.add_paragraph("Second list, same style")
for item in ("four", "five", "six"):
    doc_a.add_paragraph(item, style="List Number")
doc_a.save("continuing.docx")

# ---- document B: second list restarts via cloned numbering definition ---------
doc_b = Document()
doc_b.add_paragraph("First list")
for item in ("one", "two", "three"):
    doc_b.add_paragraph(item, style="List Number")
doc_b.add_paragraph("Second list, restarted")
base_id = list_number_num_id(doc_b)
numbering_b = doc_b.part.part_related_by(RT.NUMBERING).element
base_entry = next(n for n in numbering_b.findall(qn("w:num"))
                  if n.get(qn("w:numId")) == str(base_id))
preexisting_override = OxmlElement("w:lvlOverride")
preexisting_override.set(qn("w:ilvl"), "0")
preexisting_start = OxmlElement("w:startOverride")
preexisting_start.set(qn("w:val"), "7")
preexisting_override.append(preexisting_start)
base_entry.append(preexisting_override)
restart_id = new_restart_num_id(doc_b, base_id)
for item in ("four", "five", "six"):
    numbered_paragraph(doc_b, item, restart_id)
doc_b.save("restarted.docx")

# ---- structural assertions ------------------------------------------------------
numbering_b = doc_b.part.part_related_by(RT.NUMBERING).element
nums_b = numbering_b.findall(qn("w:num"))
check("cloned num entry exists", any(n.get(qn("w:numId")) == str(restart_id) for n in nums_b))
clone_entry = next(n for n in nums_b if n.get(qn("w:numId")) == str(restart_id))
level_zero_overrides = [item for item in clone_entry.findall(qn("w:lvlOverride"))
                        if item.get(qn("w:ilvl")) == "0"]
start_overrides = ([] if not level_zero_overrides else
                   level_zero_overrides[0].findall(qn("w:startOverride")))
check("clone replaces a preexisting ilvl=0 override without creating a duplicate",
      len(level_zero_overrides) == 1 and len(start_overrides) == 1
      and start_overrides[0].get(qn("w:val")) == "1")
check("cloning leaves the base numbering override unchanged",
      preexisting_start.get(qn("w:val")) == "7")
second_list_num_ids = [
    p._p.find(qn("w:pPr") + "/" + qn("w:numPr") + "/" + qn("w:numId")).get(qn("w:val"))
    for p in doc_b.paragraphs[-3:]
]
check("restarted paragraphs reference the cloned numId", set(second_list_num_ids) == {str(restart_id)}, second_list_num_ids)

# ---- rendered proof via LibreOffice --------------------------------------------
subprocess.run(
    ["soffice", "--headless", "--convert-to", "pdf", "--outdir", ".", "continuing.docx", "restarted.docx"],
    check=True, capture_output=True, timeout=180,
)


def second_list_numbers(pdf_path):
    text = " ".join(page.get_text() for page in fitz.open(pdf_path))
    # take the rendered numbers in front of the second list's item words
    return [text.split(word)[0].split()[-1] for word in ("four", "five", "six")]


cont = second_list_numbers("continuing.pdf")
restart = second_list_numbers("restarted.pdf")
print("continuing.docx renders second list as:", cont)
print("restarted.docx renders second list as:", restart)
check("plain style reuse continues the sequence (negative control)", cont == ["4.", "5.", "6."], cont)
check("cloned definition restarts the second list at 1", restart == ["1.", "2.", "3."], restart)

# ---- scenes.md snippet: keep_table_together for the signature block -------------
doc_c = Document()
doc_c.add_heading("Contract", level=1)
# fill most of page 1 so a tall signature table would otherwise straddle the page break
for _ in range(24):
    doc_c.add_paragraph("Filler paragraph to push the signature block toward the page break. " * 3)
sig = doc_c.add_table(rows=4, cols=2)
labels = [("甲方（盖章）", "乙方（盖章）"), ("签字", "签字"), ("日期", "日期"), ("备注", "备注")]
for r, pair in enumerate(labels):
    sig.cell(r, 0).text, sig.cell(r, 1).text = pair


def keep_table_together(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))   # a row never splits mid-row
    for row in table.rows[:-1]:
        for cell in row.cells:
            for par in cell.paragraphs:
                par.paragraph_format.keep_with_next = True  # row sticks to the next row


keep_table_together(sig)
doc_c.save("signature.docx")
reopened = Document("signature.docx")
check("every signature row carries cantSplit",
      all(row._tr.find(qn("w:trPr")) is not None and row._tr.find(qn("w:trPr")).find(qn("w:cantSplit")) is not None
          for row in reopened.tables[0].rows))
subprocess.run(
    ["soffice", "--headless", "--convert-to", "pdf", "--outdir", ".", "signature.docx"],
    check=True, capture_output=True, timeout=180,
)
pages_with_labels = [
    page.number for page in fitz.open("signature.pdf")
    if "甲方（盖章）" in page.get_text() and "备注" in page.get_text()
]
check("rendered signature table stays on one page", len(pages_with_labels) == 1, pages_with_labels)


# ---- read.md snippet: merged-cell spans survive extraction -----------------------
def table_matrix(table):
    rows = []
    for row in table.rows:
        cells = []
        for tc in row._tr.tc_lst:
            tc_pr = tc.find(qn("w:tcPr"))
            grid_span = tc_pr.find(qn("w:gridSpan")) if tc_pr is not None else None
            v_merge = tc_pr.find(qn("w:vMerge")) if tc_pr is not None else None
            note = ""
            if grid_span is not None:
                note += "(span {})".format(grid_span.get(qn("w:val")))
            if v_merge is not None:
                note += "(vmerge start)" if v_merge.get(qn("w:val")) == "restart" else "(vmerge cont.)"
            text = "".join(node.text or "" for node in tc.iter(qn("w:t")))
            cells.append(text + note if note else text)
        rows.append(cells)
    return rows

merge_doc = Document()
merge_table = merge_doc.add_table(rows=3, cols=3)
merge_table.cell(0, 0).text = "HEAD"  # set only the origin; merging concatenates texts
merge_table.cell(0, 0).merge(merge_table.cell(0, 1))
merge_table.cell(1, 2).text = "TOP"
merge_table.cell(2, 2).text = "BOTTOM"
merge_table.cell(1, 2).merge(merge_table.cell(2, 2))
merge_doc.save("merged-cells.docx")
merged_reopened = Document("merged-cells.docx")
merged_table = merged_reopened.tables[0]
naive_lengths = [len(row.cells) for row in merged_table.rows]
naive_row0 = [cell.text for cell in merged_table.rows[0].cells]
matrix = table_matrix(merged_table)
check("naive row.cells expands merges to full grid width (negative control)",
      naive_lengths == [3, 3, 3] and naive_row0[0] == naive_row0[1] == "HEAD",
      (naive_lengths, naive_row0))
check("matrix keeps one entry per real tc",
      [len(row) for row in matrix] == [2, 3, 3], matrix)
check("horizontal merge is annotated", "(span 2)" in matrix[0][0], matrix[0])
check("vertical merge start and continuation are annotated",
      any("(vmerge start)" in cell for cell in matrix[1]) and
      any("(vmerge cont.)" in cell for cell in matrix[2]), matrix)

# ---- cjk.md snippet: Hangul routes through the East Asian slot -------------------
def font_slot(character):
    codepoint = ord(character)
    return "eastAsia" if (
        0x1100 <= codepoint <= 0x11FF or 0x2E80 <= codepoint <= 0x9FFF
        or 0x3130 <= codepoint <= 0x318F or 0xA960 <= codepoint <= 0xA97F
        or 0xAC00 <= codepoint <= 0xD7FF or 0xF900 <= codepoint <= 0xFAFF
        or 0x20000 <= codepoint <= 0x3134F
        or 0x31350 <= codepoint <= 0x33479
    ) else ("ascii" if codepoint < 128 else "hAnsi")

check("Hangul syllables use the East Asian slot", font_slot("한") == "eastAsia")
check("Hangul jamo use the East Asian slot", font_slot("ᄀ") == "eastAsia")
check("compatibility jamo use the East Asian slot", font_slot("ㄱ") == "eastAsia")
check("Hangul extended-A uses the East Asian slot", font_slot(chr(0xA960)) == "eastAsia")
check("Hangul extended-B uses the East Asian slot", font_slot(chr(0xD7B0)) == "eastAsia")
check("U+2E80 CJK radical uses the East Asian slot", font_slot(chr(0x2E80)) == "eastAsia")
check("Unicode 17 supplementary Han endpoints use the East Asian slot",
      font_slot(chr(0x31350)) == font_slot(chr(0x33479)) == "eastAsia")
check("the codepoint after Unicode 17 Han remains hAnsi",
      font_slot(chr(0x3347A)) == "hAnsi")
check("Latin stays in the ascii slot", font_slot("A") == "ascii")
check("non-CJK fullwidth-range-adjacent Latin-1 stays hAnsi", font_slot("é") == "hAnsi")



# ---- read.md tc_text: paragraph boundaries survive cell extraction ---------------
def tc_text(tc):
    paragraphs = []
    for p in tc.iter(qn("w:p")):
        pieces = []
        for node in p.iter():
            if node.tag == qn("w:t"):
                pieces.append(node.text or "")
            elif node.tag == qn("w:tab"):
                pieces.append("<tab>")
            elif node.tag in (qn("w:br"), qn("w:cr")):
                pieces.append("<br>")
            elif node.tag == qn("w:noBreakHyphen"):
                pieces.append("-")
        paragraphs.append("".join(pieces))
    return " / ".join(paragraphs)

cells_doc = Document()
cells_table = cells_doc.add_table(rows=1, cols=1)
cell = cells_table.cell(0, 0)
cell.paragraphs[0].text = "First"
cell.add_paragraph("Second")
run_with_tab = cell.paragraphs[0].add_run("")
tab_element = OxmlElement("w:tab")
run_with_tab._r.append(tab_element)
run_with_tab.add_text("after tab")
hyphen_run = cell.paragraphs[0].add_run("non")
hyphen_run._r.append(OxmlElement("w:noBreakHyphen"))
hyphen_run.add_text("breaking")
cells_doc.save("cell-paragraphs.docx")
cells_reopened = Document("cell-paragraphs.docx")
cells_tc = cells_reopened.tables[0].rows[0]._tr.tc_lst[0]
joined_raw = "".join(node.text or "" for node in cells_tc.iter(qn("w:t")))
extracted = tc_text(cells_tc)
check("raw w:t joining concatenates paragraphs (negative control)",
      "Firstafter tab" in joined_raw, joined_raw)
check("tc_text preserves the paragraph boundary", " / Second" in extracted, extracted)
check("tc_text keeps tabs visible", "<tab>after tab" in extracted, extracted)
check("tc_text preserves a table-cell nonbreaking hyphen",
      "non-breaking" in extracted, extracted)


# ---- edit.md raw OOXML repack: every input gets a fresh extraction tree ---------
media_doc = Document()
media_doc.add_paragraph("first document")
media_doc.add_picture("inline-icon.png")
media_doc.save("repack-with-media.docx")
plain_doc = Document()
plain_doc.add_paragraph("second document")
plain_doc.save("repack-without-media.docx")


WINDOWS_DEVICE_NAMES = {
    "con", "prn", "aux", "nul", "conin$", "conout$",
    *(f"com{suffix}" for suffix in "123456789¹²³"),
    *(f"lpt{suffix}" for suffix in "123456789¹²³"),
}


def extraction_key(name):
    is_directory = name.endswith("/")
    path = name[:-1] if is_directory else name
    require(path and not name.startswith("/") and "\\" not in name,
            f"non-canonical archive member path: {name}")
    parts = path.split("/")
    require(len(parts) <= MAX_MEMBER_COMPONENTS,
            "archive member depth exceeds portable extraction limit")
    require(
        all(
            part not in {"", ".", ".."}
            and not any(character in ':<>|"?*' for character in part)
            and not any(ord(character) < 32 for character in part)
            and not part.endswith((".", " "))
            and unicodedata.normalize("NFC", part) == part
            and part.partition(".")[0].rstrip(" ").casefold() not in WINDOWS_DEVICE_NAMES
            for part in parts
        ),
        f"non-canonical archive member path: {name}",
    )
    require(
        all(
            len(part.encode("utf-8")) <= MAX_MEMBER_COMPONENT_BYTES
            and len(part.encode("utf-16-le")) // 2 <= MAX_MEMBER_COMPONENT_UTF16_UNITS
            for part in parts
        ),
        "archive member component exceeds portable extraction limit",
    )
    require(len(path.encode("utf-8")) <= MAX_MEMBER_PATH_BYTES
            and len(path.encode("utf-16-le")) // 2 <= MAX_MEMBER_PATH_UTF16_UNITS,
            "archive member path exceeds portable extraction limit")
    canonical = PurePosixPath(*parts).as_posix()
    require(canonical == path, f"non-canonical archive member path: {name}")
    key = tuple(unicodedata.normalize("NFC", part.casefold()) for part in parts)
    return key, tuple(parts), is_directory


def validate_extraction_paths(infos):
    root = {"children": {}, "member": False, "file": False, "spelling": None}
    for info in infos:
        key, spellings, is_directory = extraction_key(info.filename)
        node = root
        for normalized, spelling in zip(key, spellings):
            require(not node["file"],
                    "archive file and directory paths collide after extraction")
            child = node["children"].get(normalized)
            if child is None:
                child = {
                    "children": {}, "member": False, "file": False,
                    "spelling": spelling,
                }
                node["children"][normalized] = child
            else:
                require(child["spelling"] == spelling,
                        "archive member path spelling collides after extraction")
            node = child
        require(not node["member"], "archive member paths collide after extraction")
        require(is_directory or not node["children"],
                "archive file and directory paths collide after extraction")
        node["member"] = True
        node["file"] = not is_directory


def validate_docx_archive_bounds(archive):
    require(os.fstat(archive.fp.fileno()).st_size <= MAX_ARCHIVE_BYTES,
            "compressed DOCX file size above limit")
    infos = archive.infolist()
    require(len(infos) <= MAX_MEMBERS, "archive member count above limit")
    names = {info.filename for info in infos}
    require(len(names) == len(infos), "duplicate archive member names are unsafe")
    validate_extraction_paths(infos)
    require("[Content_Types].xml" in names and "word/document.xml" in names,
            "required DOCX package parts are missing")
    require(sum(info.file_size for info in infos) <= MAX_TOTAL_UNCOMPRESSED,
            "declared archive size exceeds the edit limit")
    actual_total = 0
    for info in infos:
        require(info.file_size <= MAX_ENTRY, f"oversized part: {info.filename}")
        require(info.file_size / max(info.compress_size, 1) <= MAX_COMPRESSION_RATIO,
                f"suspicious compression ratio: {info.filename}")
        if info.filename.endswith((".xml", ".rels")):
            require(info.file_size <= MAX_XML_PART, f"oversized XML part: {info.filename}")
        actual_size = 0
        with archive.open(info) as stream:
            while chunk := stream.read(64 * 1024):
                actual_size += len(chunk)
                actual_total += len(chunk)
                require(actual_size <= MAX_ENTRY, f"part exceeded read limit: {info.filename}")
                require(actual_total <= MAX_TOTAL_UNCOMPRESSED,
                        "archive exceeded total read limit")
        require(actual_size == info.file_size, f"size mismatch: {info.filename}")


def repack_tree(source, output, extraction_root):
    with zipfile.ZipFile(source) as archive:
        validate_docx_archive_bounds(archive)
        archive.extractall(extraction_root)
    content_types = Path(extraction_root) / "[Content_Types].xml"
    files = sorted(
        (path for path in Path(extraction_root).rglob("*")
         if path.is_file() and path != content_types),
        key=lambda path: path.relative_to(extraction_root).as_posix(),
    )
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED, strict_timestamps=False) as archive:
        archive.write(content_types, "[Content_Types].xml")
        for path in files:
            archive.write(path, path.relative_to(extraction_root).as_posix())


bomb_extract_root = Path("bomb-extract")
bomb_extract_root.mkdir()
try:
    repack_tree("compressed-bomb.docx", "bomb-output.docx", bomb_extract_root)
    edit_bomb_rejected_before_extract = False
except ValueError as error:
    edit_bomb_rejected_before_extract = (
        str(error) == "suspicious compression ratio: word/document.xml"
        and
        not any(bomb_extract_root.iterdir()) and not Path("bomb-output.docx").exists()
    )
check("Tier 2 edit rejects an archive bomb before extracting any member",
      edit_bomb_rejected_before_extract)
check("Tier 2 pre-extract bounds remain active under optimized Python",
      __debug__ or edit_bomb_rejected_before_extract)

with zipfile.ZipFile("noncanonical-member.docx", "w", zipfile.ZIP_STORED) as archive:
    archive.writestr("[Content_Types].xml", "<Types/>")
    archive.writestr("word/document.xml", "<canonical/>")
    archive.writestr("../word/document.xml", "<shadow/>")
with TemporaryDirectory(prefix="docx-noncanonical-") as scratch:
    noncanonical_root = Path(scratch)
    try:
        repack_tree("noncanonical-member.docx", "noncanonical-output.docx",
                    noncanonical_root)
        noncanonical_rejected = False
    except ValueError as error:
        noncanonical_rejected = (
            str(error) == "non-canonical archive member path: ../word/document.xml"
            and not any(noncanonical_root.iterdir())
            and not Path("noncanonical-output.docx").exists()
        )
check("Tier 2 rejects traversal aliases before they can overwrite a validated part",
      noncanonical_rejected)
check("non-canonical member rejection remains active under optimized Python",
      __debug__ or noncanonical_rejected)

with zipfile.ZipFile("normalized-collision.docx", "w", zipfile.ZIP_STORED) as archive:
    archive.writestr("[Content_Types].xml", "<Types/>")
    archive.writestr("word/document.xml", "<document/>")
    archive.writestr("custom", b"file")
    archive.writestr("custom/", b"")
with TemporaryDirectory(prefix="docx-collision-") as scratch:
    collision_root = Path(scratch)
    try:
        repack_tree("normalized-collision.docx", "collision-output.docx", collision_root)
        normalized_collision_rejected = False
    except ValueError as error:
        normalized_collision_rejected = (
            str(error) == "archive member paths collide after extraction"
            and not any(collision_root.iterdir())
            and not Path("collision-output.docx").exists()
        )
check("Tier 2 rejects distinct names that normalize to one extraction path",
      normalized_collision_rejected)

with zipfile.ZipFile("prefix-collision.docx", "w", zipfile.ZIP_STORED) as archive:
    archive.writestr("[Content_Types].xml", "<Types/>")
    archive.writestr("word/document.xml", "<document/>")
    archive.writestr("custom", b"file")
    archive.writestr("custom/child.bin", b"child")
with TemporaryDirectory(prefix="docx-prefix-collision-") as scratch:
    prefix_collision_root = Path(scratch)
    try:
        repack_tree("prefix-collision.docx", "prefix-collision-output.docx",
                    prefix_collision_root)
        prefix_collision_rejected = False
    except ValueError as error:
        prefix_collision_rejected = (
            str(error) == "archive file and directory paths collide after extraction"
            and not any(prefix_collision_root.iterdir())
            and not Path("prefix-collision-output.docx").exists()
        )
check("Tier 2 rejects a file path that is also an extracted directory prefix",
      prefix_collision_rejected)

with zipfile.ZipFile("case-collision.docx", "w", zipfile.ZIP_STORED) as archive:
    archive.writestr("[Content_Types].xml", "<Types/>")
    archive.writestr("word/document.xml", "<canonical/>")
    archive.writestr("WORD/document.xml", "<shadow/>")
with TemporaryDirectory(prefix="docx-case-collision-") as scratch:
    case_collision_root = Path(scratch)
    try:
        repack_tree("case-collision.docx", "case-collision-output.docx",
                    case_collision_root)
        case_collision_rejected = False
    except ValueError as error:
        case_collision_rejected = (
            str(error) == "archive member path spelling collides after extraction"
            and not any(case_collision_root.iterdir())
            and not Path("case-collision-output.docx").exists()
        )
check("Tier 2 rejects case aliases before Windows extraction can overwrite a part",
      case_collision_rejected)

with zipfile.ZipFile("prefix-spelling-collision.docx", "w", zipfile.ZIP_STORED) as archive:
    archive.writestr("[Content_Types].xml", "<Types/>")
    archive.writestr("word/document.xml", "<document/>")
    archive.writestr("WORD/styles.xml", "<styles/>")
with TemporaryDirectory(prefix="docx-prefix-spelling-") as scratch:
    prefix_spelling_root = Path(scratch)
    try:
        repack_tree("prefix-spelling-collision.docx", "prefix-spelling-output.docx",
                    prefix_spelling_root)
        prefix_spelling_rejected = False
    except ValueError as error:
        prefix_spelling_rejected = (
            str(error) == "archive member path spelling collides after extraction"
            and not any(prefix_spelling_root.iterdir())
            and not Path("prefix-spelling-output.docx").exists()
        )
check("Tier 2 rejects case aliases in a shared directory prefix",
      prefix_spelling_rejected)

with zipfile.ZipFile("overlong-member.docx", "w", zipfile.ZIP_STORED) as archive:
    archive.writestr("[Content_Types].xml", "<Types/>")
    archive.writestr("word/document.xml", "<document/>")
    archive.writestr("word/" + "a" * 256, b"oversized component")
with TemporaryDirectory(prefix="docx-overlong-member-") as scratch:
    overlong_root = Path(scratch)
    try:
        repack_tree("overlong-member.docx", "overlong-output.docx", overlong_root)
        overlong_member_rejected = False
    except ValueError as error:
        overlong_member_rejected = (
            str(error) == "archive member component exceeds portable extraction limit"
            and not any(overlong_root.iterdir())
            and not Path("overlong-output.docx").exists()
        )
check("Tier 2 rejects an unportable component before partial Windows extraction",
      overlong_member_rejected)

portable_path_limits_rejected = []
for unsafe_name in (
    "/".join(["a"] * (MAX_MEMBER_COMPONENTS + 1)),
    "/".join(["a" * 80] * 4),
):
    try:
        extraction_key(unsafe_name)
    except ValueError:
        portable_path_limits_rejected.append(unsafe_name)
check("portable extraction bounds total path length and component depth",
      len(portable_path_limits_rejected) == 2)

portable_name_rejections = []
for unsafe_name in (
    "/word/document.xml", "word\\document.xml", "word/con.xml", "word/COM¹.xml",
    "word/LPT².txt", "word/CONIN$.xml", "word/CONOUT$.xml", "word/NUL .xml",
    "word/trailing. ", "word/cafe\u0301.xml", "word/control\x01.xml",
):
    try:
        extraction_key(unsafe_name)
    except ValueError:
        portable_name_rejections.append(unsafe_name)
check("portable extraction rejects absolute, alternate, device, non-NFC, and control names",
      len(portable_name_rejections) == 11, portable_name_rejections)

with zipfile.ZipFile("malformed-for-repair.docx", "w", zipfile.ZIP_STORED) as archive:
    archive.writestr("[Content_Types].xml", "<Types/>")
    archive.writestr("word/document.xml", "<w:document")
with TemporaryDirectory(prefix="docx-malformed-repair-") as scratch:
    malformed_root = Path(scratch)
    repack_tree("malformed-for-repair.docx", "malformed-repacked.docx", malformed_root)
    malformed_extracted = (malformed_root / "word" / "document.xml").is_file()
check("archive bounds allow Tier 2 to extract bounded malformed XML for repair",
      malformed_extracted)

reused_root = Path("reused-docx-work")
reused_root.mkdir()
repack_tree("repack-with-media.docx", "reused-first.docx", reused_root)
repack_tree("repack-without-media.docx", "reused-second.docx", reused_root)
with zipfile.ZipFile("reused-second.docx") as archive:
    reused_names = archive.namelist()
check("reusing an extraction tree demonstrably leaks a prior document's media",
      any(name.startswith("word/media/") for name in reused_names), reused_names)

with TemporaryDirectory(prefix="docx-edit-") as scratch:
    repack_tree("repack-without-media.docx", "fresh-second.docx", Path(scratch))
with zipfile.ZipFile("fresh-second.docx") as archive:
    fresh_names = archive.namelist()
check("fresh extraction excludes media absent from the current input",
      not any(name.startswith("word/media/") for name in fresh_names), fresh_names)
check("fresh repack writes Content_Types exactly once and first",
      fresh_names[0] == "[Content_Types].xml"
      and fresh_names.count("[Content_Types].xml") == 1, fresh_names[:3])
check("freshly repacked DOCX reopens with only the current document's content",
      Document("fresh-second.docx").paragraphs[0].text == "second document")


print("\n" + ("ALL DOCX FIXTURES PASSED" if not failures else f"{len(failures)} FAILURES: {failures}"))
sys.exit(0 if not failures else 1)
